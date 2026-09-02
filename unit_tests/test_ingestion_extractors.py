"""Unit tests for legal_rag.ingestion.extractors.

Scope
-----
- TxtExtractor / DocxExtractor: exercised against real tiny files (fast,
  no heavy mocking needed -- these libraries are lightweight).
- PdfExtractor.extract_page(): the page-level decision logic (native vs.
  OCR vs. failed) is tested directly with a fake PDF page and a fake OCR
  engine, which is what actually needs unit-test coverage; nobody should
  need a real scanned legal PDF to prove this branch logic works.
- PdfExtractor.extract(): exercised end-to-end against a real, tiny PDF
  built at test time with pymupdf itself.
"""

from __future__ import annotations

import pymupdf
import pytest
from docx import Document as DocxDocument

from legal_rag.ingestion.extractors import DocxExtractor, PdfExtractor, TxtExtractor
from legal_rag.ingestion.models import (
    DocumentMetadata,
    ExtractionMethod,
    OcrText,
    SourceFormat,
    TextQuality,
)


def _metadata(source_format=SourceFormat.TXT) -> DocumentMetadata:
    return DocumentMetadata(
        document_id="doc-1",
        document_version=1,
        document_type="statute",
        source="unit-tests",
        source_file="doc",
        file_hash="hash",
        source_format=source_format,
    )


# --------------------------------------------------------------- TxtExtractor


def test_txt_extractor_yields_one_record_with_line_locators(tmp_path):
    path = tmp_path / "doc.txt"
    path.write_text("Article 1: text\nArticle 2: more text\n", encoding="utf-8")

    extractor = TxtExtractor(expected_language="ar")
    records = list(extractor.extract(path, _metadata()))

    assert len(records) == 1
    record = records[0]
    assert record.extraction_method is ExtractionMethod.TXT
    assert "Article 1" in record.original_text
    assert record.locator_start == 1
    assert record.locator_end >= 1


def test_txt_extractor_rejects_page_limit():
    extractor = TxtExtractor(expected_language="ar")
    with pytest.raises(ValueError, match="only for PDF input"):
        list(extractor.extract(None, _metadata(), page_limit=1))


# -------------------------------------------------------------- DocxExtractor


def test_docx_extractor_reads_paragraphs_and_tables(tmp_path):
    path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_paragraph("Article 1: Introductory text.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Definition"
    document.save(str(path))

    extractor = DocxExtractor(expected_language="en")
    records = list(extractor.extract(path, _metadata(SourceFormat.DOCX)))

    assert len(records) == 1
    assert "Article 1" in records[0].original_text
    assert "Term" in records[0].original_text
    assert records[0].extraction_method is ExtractionMethod.DOCX


def test_docx_extractor_rejects_empty_document(tmp_path):
    path = tmp_path / "empty.docx"
    document = DocxDocument()
    document.save(str(path))

    extractor = DocxExtractor(expected_language="en")
    with pytest.raises(ValueError, match="no extractable body text"):
        list(extractor.extract(path, _metadata(SourceFormat.DOCX)))


def test_docx_extractor_rejects_page_limit():
    extractor = DocxExtractor(expected_language="en")
    with pytest.raises(ValueError, match="only for PDF input"):
        list(extractor.extract(None, _metadata(SourceFormat.DOCX), page_limit=1))


# --------------------------------------------------------------- PdfExtractor


def _pdf_metadata() -> DocumentMetadata:
    return _metadata(SourceFormat.PDF)


class _FakeOcrEngine:
    def __init__(self, text: str, confidence: float | None = 0.9):
        self._text = text
        self._confidence = confidence

    def extract_page(self, page):
        return OcrText(text=self._text, mean_confidence=self._confidence)


class _FakePdfPage:
    """Minimal stand-in for a pymupdf page as used by extract_page()."""

    def get_text(self, mode, sort=False):
        return {"blocks": []}  # no RTL digit runs detected


def test_extract_page_uses_native_text_when_quality_is_sufficient():
    good_arabic_text = "نص قانوني عربي واضح وطويل بما يكفي لتجاوز حد الجودة. " * 3
    extractor = PdfExtractor(
        expected_language="ar",
        ocr_engine_factory=lambda language: _FakeOcrEngine("should not be used"),
    )

    record = extractor.extract_page(
        _FakePdfPage(), native_text=good_arabic_text, metadata=_pdf_metadata(), page_number=1
    )

    assert record.extraction_method is ExtractionMethod.NATIVE
    assert record.original_text == good_arabic_text


def test_extract_page_falls_back_to_ocr_for_low_quality_native_text():
    poor_text = "x"  # far too short -> requires_ocr() is True
    good_ocr_text = "نص عربي مستخرج بواسطة أوسي آر بجودة كافية للمرور. " * 3

    extractor = PdfExtractor(
        expected_language="ar",
        ocr_engine_factory=lambda language: _FakeOcrEngine(good_ocr_text),
    )

    record = extractor.extract_page(
        _FakePdfPage(), native_text=poor_text, metadata=_pdf_metadata(), page_number=2
    )

    assert record.extraction_method is ExtractionMethod.OCR
    assert record.original_text == good_ocr_text
    assert record.page_number == 2


def test_extract_page_marks_failed_when_ocr_output_is_also_poor():
    poor_text = "x"
    poor_ocr_text = "x"  # also fails the quality check

    extractor = PdfExtractor(
        expected_language="ar",
        ocr_engine_factory=lambda language: _FakeOcrEngine(poor_ocr_text),
    )

    record = extractor.extract_page(
        _FakePdfPage(), native_text=poor_text, metadata=_pdf_metadata(), page_number=3
    )

    assert record.extraction_method is ExtractionMethod.FAILED
    assert record.error is not None
    assert record.original_text == ""


def test_extract_page_auto_language_tries_both_ocr_languages():
    poor_text = "1"
    good_english_text = "This is clear English OCR output that is long enough to pass. " * 2

    calls: list[str] = []

    def factory(language):
        calls.append(language)
        if language == "en":
            return _FakeOcrEngine(good_english_text)
        return _FakeOcrEngine("x")  # Arabic OCR fails quality

    extractor = PdfExtractor(expected_language="auto", ocr_engine_factory=factory)

    record = extractor.extract_page(
        _FakePdfPage(), native_text=poor_text, metadata=_pdf_metadata(), page_number=1
    )

    assert record.extraction_method is ExtractionMethod.OCR
    assert "en" in calls


def test_extract_end_to_end_against_a_real_tiny_pdf(tmp_path):
    """Build a real one-page PDF and run it through the real extract()
    pipeline (real pymupdf + real pypdf). Depending on how pypdf's text
    layer extraction handles this minimal, programmatically generated
    PDF, the native text may or may not pass the quality bar -- so a
    reasonable OCR fallback text is supplied and either successful path
    (NATIVE or OCR) is accepted as correct. What matters for this unit
    test is that extract() runs the real end-to-end flow without errors
    and returns exactly one well-formed, non-failed record.
    """

    path = tmp_path / "doc.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Article 1: This is a real embedded PDF text layer.")
    document.save(str(path))
    document.close()

    fallback_ocr_text = "Article 1: OCR fallback text long enough to pass quality checks. " * 2
    extractor = PdfExtractor(
        expected_language="en",
        ocr_engine_factory=lambda language: _FakeOcrEngine(fallback_ocr_text),
    )
    records = list(extractor.extract(path, _pdf_metadata()))

    assert len(records) == 1
    assert records[0].page_number == 1
    assert records[0].extraction_method in (ExtractionMethod.NATIVE, ExtractionMethod.OCR)
    assert "Article 1" in records[0].original_text
