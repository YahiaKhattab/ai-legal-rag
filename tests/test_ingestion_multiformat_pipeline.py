import json
from pathlib import Path

from docx import Document

from legal_rag.ingestion.models import IngestionStatus
from legal_rag.ingestion.pipeline import IngestionPipeline


class WordTokenCounter:
    name = "test-word-counter"

    def count_passage(self, text: str) -> int:
        return len(text.split()) + 3

    def count_content(self, text: str) -> int:
        return len(text.split())


def _pipeline() -> IngestionPipeline:
    return IngestionPipeline(expected_language="ar", token_counter=WordTokenCounter())


def _chunks(path: Path) -> list[dict[str, object]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines()]


def test_txt_pipeline_writes_line_citations_and_reuses_duplicate(tmp_path: Path) -> None:
    path = tmp_path / "law.txt"
    path.write_text("المادة الأولى\n" + "نص قانوني مفهوم " * 40, encoding="utf-8")
    output = tmp_path / "processed"

    first = _pipeline().ingest(path, output, document_type="law", source="Test Authority")
    second = _pipeline().ingest(path, output, document_type="law", source="Test Authority")

    assert first.status is IngestionStatus.PROCESSED
    assert second.status is IngestionStatus.DUPLICATE
    assert first.sources_output.name.endswith(".sources.jsonl")
    [source] = [
        json.loads(line) for line in first.sources_output.read_text(encoding="utf-8").splitlines()
    ]
    assert source["source_format"] == "txt"
    assert source["locator_type"] == "line"
    chunks = _chunks(first.chunks_output)
    assert chunks
    assert all(chunk["page_start"] is None for chunk in chunks)
    assert all(chunk["locator_type"] == "line" for chunk in chunks)
    assert all(
        isinstance(chunk["locator_start"], int) and chunk["locator_start"] >= 1 for chunk in chunks
    )


def test_docx_pipeline_writes_block_citations(tmp_path: Path) -> None:
    path = tmp_path / "law.docx"
    document = Document()
    document.add_paragraph("المادة الأولى")
    document.add_paragraph("نص قانوني مفهوم " * 20)
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "المصطلح"
    table.cell(0, 1).text = "التعريف"
    document.save(str(path))

    summary = _pipeline().ingest(path, tmp_path / "processed")
    chunks = _chunks(summary.chunks_output)

    assert summary.direct_records == 1
    assert summary.ocr_records == 0
    assert chunks
    assert all(chunk["source_format"] == "docx" for chunk in chunks)
    assert all(chunk["locator_type"] == "block" for chunk in chunks)
    assert all(chunk["page_start"] is None for chunk in chunks)


def test_auto_language_assigns_language_per_bilingual_chunk(tmp_path: Path) -> None:
    path = tmp_path / "bilingual-law.txt"
    path.write_text(
        "المادة (1)\n" + "هذا نص قانوني عربي واضح يتعلق بالخدمات المصرفية. " * 8 + "\n"
        "Article 2\n" + "This is clear English legal text about regulated banking services. " * 8,
        encoding="utf-8",
    )
    pipeline = IngestionPipeline(expected_language="auto", token_counter=WordTokenCounter())

    summary = pipeline.ingest(path, tmp_path / "processed")
    [source] = [
        json.loads(line) for line in summary.sources_output.read_text(encoding="utf-8").splitlines()
    ]
    chunks = _chunks(summary.chunks_output)

    assert source["language"] == "mixed"
    assert [chunk["language"] for chunk in chunks] == ["ar", "en"]
