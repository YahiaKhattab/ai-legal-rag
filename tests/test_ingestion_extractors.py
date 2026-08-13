# ruff: noqa: RUF001 - Arabic fixtures intentionally exercise legal text

from pathlib import Path
from typing import Any

import pytest
from docx import Document

from legal_rag.ingestion.extractors import DocxExtractor, OcrLanguage, PdfExtractor, TxtExtractor
from legal_rag.ingestion.models import (
    DocumentMetadata,
    ExtractionMethod,
    LocatorType,
    OcrText,
    SourceFormat,
)


def _metadata(path: Path, source_format: SourceFormat) -> DocumentMetadata:
    return DocumentMetadata(
        document_id="a" * 64,
        document_version=1,
        document_type="law",
        source="Test Authority",
        source_file=path.name,
        file_hash="a" * 64,
        source_format=source_format,
    )


def test_txt_extractor_preserves_text_and_line_locators(tmp_path: Path) -> None:
    path = tmp_path / "law.txt"
    original = "المادة الأولى\nالنص الأول\nالنص الثاني"
    path.write_bytes(original.encode("utf-8"))

    [record] = list(
        TxtExtractor(expected_language="ar").extract(
            path,
            _metadata(path, SourceFormat.TXT),
        )
    )

    assert record.extraction_method is ExtractionMethod.TXT
    assert record.original_text == original
    assert record.locator_type is LocatorType.LINE
    assert record.locator_start == 1
    assert record.locator_end == 3
    assert [segment.locator_start for segment in record.source_segments] == [1, 2, 3]


def test_docx_extractor_preserves_paragraph_and_table_order(tmp_path: Path) -> None:
    path = tmp_path / "law.docx"
    document = Document()
    document.add_paragraph("المادة الأولى")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "المصطلح"
    table.cell(0, 1).text = "التعريف"
    document.add_paragraph("النص الختامي")
    document.save(str(path))

    [record] = list(
        DocxExtractor(expected_language="ar").extract(
            path,
            _metadata(path, SourceFormat.DOCX),
        )
    )

    assert record.extraction_method is ExtractionMethod.DOCX
    assert record.locator_type is LocatorType.BLOCK
    assert record.original_text.splitlines() == [
        "المادة الأولى",
        "المصطلح\tالتعريف",
        "النص الختامي",
    ]
    assert [segment.kind for segment in record.source_segments] == [
        "paragraph",
        "table_row",
        "paragraph",
    ]


def test_non_pdf_extractors_reject_page_limit(tmp_path: Path) -> None:
    path = tmp_path / "law.txt"
    path.write_text("legal text", encoding="utf-8")

    with pytest.raises(ValueError, match="only for PDF"):
        list(
            TxtExtractor(expected_language="en").extract(
                path,
                _metadata(path, SourceFormat.TXT),
                page_limit=1,
            )
        )


def test_txt_auto_language_marks_bilingual_content_as_mixed(tmp_path: Path) -> None:
    path = tmp_path / "bilingual-law.txt"
    path.write_text(
        "المادة الأولى من القانون المصرفي\nArticle one of the banking regulation",
        encoding="utf-8",
    )

    [record] = list(
        TxtExtractor(expected_language="auto").extract(
            path,
            _metadata(path, SourceFormat.TXT),
        )
    )

    assert record.language == "mixed"


def test_pdf_auto_ocr_probes_both_languages_then_reuses_winner(tmp_path: Path) -> None:
    path = tmp_path / "scanned.pdf"
    calls: list[str] = []

    class StaticOcrEngine:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_page(self, page: Any) -> OcrText:
            del page
            return OcrText(text=self._text, mean_confidence=0.9)

    engines = {
        "ar": StaticOcrEngine("نص عربي قانوني واضح " * 8),
        "en": StaticOcrEngine("Recognized English legal text " * 8),
    }

    def engine_factory(language: OcrLanguage) -> StaticOcrEngine:
        calls.append(language)
        return engines[language]

    extractor = PdfExtractor(
        expected_language="auto",
        ocr_engine_factory=engine_factory,
    )
    metadata = _metadata(path, SourceFormat.PDF)
    first = extractor.extract_page(object(), native_text="", metadata=metadata, page_number=1)
    calls_after_first = list(calls)
    second = extractor.extract_page(object(), native_text="", metadata=metadata, page_number=2)

    assert calls_after_first == ["ar", "en"]
    assert calls == ["ar", "en", "en"]
    assert first.extraction_method is ExtractionMethod.OCR
    assert first.language == "en"
    assert second.language == "en"
